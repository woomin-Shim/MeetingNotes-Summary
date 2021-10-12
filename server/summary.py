import docx 
from typing import List
from textrankr import TextRank
from docx.oxml.ns import qn
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

#문서 요약 부분
class MyTokenizer:
    def __call__(self, text: str) -> List[str]:
        tokens: List[str] = text.split()
        return tokens
#전체 문장을 담을 문자열
global text

#파일 열기
f = open("C:/text/text2.txt", 'r', 500,"utf-8")
text = f.read()
f.close()

mytokenizer: MyTokenizer = MyTokenizer()
textrank: TextRank = TextRank(mytokenizer)

#요약할 줄 수
line_num: int = 10

summarized: str = textrank.summarize(text, line_num)

#워드 작성 부분
doc = docx.Document()
doc.add_heading('회의록', level = 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
#para = doc.add_paragraph()
#run = para.add_run(summarized)


table = doc.add_table(rows = 4, cols = 4)
table.style = doc.styles['Table Grid']
table.columns[0].width = docx.shared.Cm(20.0) 
hdr_cell1 = table.rows[0].cells
hdr_cell1[0].width = docx.shared.Cm(20)
hdr_cell1[0].paragraphs[0].add_run('회의일시').bold = True
hdr_cell1[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  #가운데 정렬
hdr_cell1[1].paragraphs[0].add_run('').bold = True
hdr_cell1[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr_cell1[2].paragraphs[0].add_run('작성자').bold = True
hdr_cell1[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr_cell1[3].paragraphs[0].add_run('').bold = True
hdr_cell1[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

hdr_cell2 = table.rows[1].cells
hdr_cell2[0].paragraphs[0].add_run('참석자').bold = True
hdr_cell2[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  #가운데 정렬
hdr_cell2[1].merge(hdr_cell2[-1])

hdr_cell3 = table.rows[2].cells
hdr_cell3[0].paragraphs[0].add_run('회의제목').bold = True
hdr_cell3[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  #가운데 정렬
hdr_cell3[1].merge(hdr_cell3[-1])


hdr_cell4 = table.rows[3].cells
hdr_cell4[0].paragraphs[0].add_run('회의내용').bold = True
hdr_cell4[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  #가운데 정렬
hdr_cell4[1].merge(hdr_cell4[-1])
hdr_cell4[1].paragraphs[0].add_run(summarized).bold = True

#저장할 이름
doc.save('text.docx')
